# filepath: c:\Users\cruiube.UPVNET\Downloads\Carlos\Python Scripts\UPE_INF_Script_Master\script-launcher-app\scripts\blog\BlogOfertasWordPrep.py

import os
import re
import sys
from datetime import datetime
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import tkinter as tk
from tkinterdnd2 import DND_FILES, TkinterDnD
from tkinter import messagebox, filedialog

def extract_text_from_pdf(pdf_path):
    reader = PdfReader(pdf_path)
    text = ''
    for page in reader.pages:
        text += page.extract_text()
    return text

def extract_offer_details(text):
    cargo = None
    referencia = None
    link = None

    lineas = text.split('\n')

    # Extract cargo
    for i, linea in enumerate(lineas):
        if 'Puesto' in linea and not 'Sie' in linea and not 'estable' in linea:
            cargo = lineas[i].strip().replace('Puesto', '') if i + 1 < len(lineas) else None
            cargo = cargo.lower().title() if cargo else None

    # Extract referencia
    referencia_match = re.search(r'Referencia+(DL-\d{5}|E-\d{5}-UPV)', text)
    if referencia_match:
        referencia = referencia_match.group(1)
        if referencia.startswith('DL-'):
            link = f"https://aplicat.upv.es/dire-app/verOferta.xhtml?idOferta={referencia[3:]}"
        elif referencia.startswith('E-'):
            link = f"https://aplicat.upv.es/dire-app/verOffertaInt.xhtml?idOferta={referencia[2:7]}&ambito=UPV&tipo=E"
    else:
        print("No se encontró la referencia en el texto.")

    return cargo, referencia, link

# Carpeta de descargas del usuario por defecto
default_output_dir = os.path.join(os.environ["USERPROFILE"], "Downloads")
output_dir_var = None  # Se crea después de root

def create_word_document(offers):
    today = datetime.today().strftime('%d_%m_%Y')
    doc = Document()
    doc.add_heading('Ofertas de Trabajo', 0)

    for offer in offers:
        cargo, referencia, link = offer
        if cargo and referencia and link:
            p = doc.add_paragraph()
            run = p.add_run(f"{cargo}: ({referencia}) ")
            run.bold = True
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), doc.part.relate_to(link, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True))
            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rStyle = OxmlElement('w:rStyle')
            rStyle.set(qn('w:val'), 'Hyperlink')
            rPr.append(rStyle)
            new_run.append(rPr)
            new_run.text = "Link a la oferta"
            hyperlink.append(new_run)
            p._element.append(hyperlink)
        else:
            print(f"Detalles incompletos para la oferta: {offer}")

    output_dir = r"C:\Users\cruiube.UPVNET\Downloads\Carlos\Ofertas Blog"
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f"OfertasBlog_{today}.docx")
    doc.save(output_path)
    print(f"Documento guardado en: {output_path}")

    output_dir = output_dir_var.get()
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f"OfertasBlog_{today}.docx")
    doc.save(output_path)
    print(f"Documento guardado en: {output_path}")

def main(pdf_paths):
    offers = []
    for pdf_path in pdf_paths:
        text = extract_text_from_pdf(pdf_path)
        print(f"Texto extraído del PDF {pdf_path}:\n{text}\n")
        offer_details = extract_offer_details(text)
        offers.append(offer_details)
    create_word_document(offers)

def on_drop(event):
    pdf_paths = root.tk.splitlist(event.data)
    main(pdf_paths)
    messagebox.showinfo("Información", "Documento de Word generado con éxito.")

def seleccionar_carpeta():
    carpeta = filedialog.askdirectory()
    if carpeta:
        output_dir_var.set(carpeta)

if __name__ == "__main__":
    root = TkinterDnD.Tk()
    root.iconbitmap("./assets/pdf_word.ico")
    root.title("Arrastra y suelta los archivos PDF aquí")
    root.geometry("450x300")

    output_dir_var = tk.StringVar(root)
    output_dir_var.set(default_output_dir)

    frame = tk.Frame(root)
    frame.pack(fill="x", padx=10, pady=10)

    tk.Label(frame, text="Carpeta de guardado:", font=("Arial", 10)).pack(side="left")
    tk.Entry(frame, textvariable=output_dir_var, width=35).pack(side="left", padx=5)
    tk.Button(frame, text="Seleccionar...", command=seleccionar_carpeta).pack(side="left")

    label = tk.Label(root, text="Arrastra y suelta los archivos PDF aquí", font=("Arial", 12), padx=10, pady=10, borderwidth=2, relief="solid")
    label.pack(expand=True, fill="both", padx=10, pady=10)

    root.drop_target_register(DND_FILES)
    root.dnd_bind('<<Drop>>', on_drop)

    root.mainloop()